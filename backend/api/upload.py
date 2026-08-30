"""文件上传与解析接口。"""

from __future__ import annotations

from io import BytesIO
from typing import List

from fastapi import APIRouter, File, UploadFile

router = APIRouter(prefix="/upload", tags=["Upload / 文件上传"])


@router.post("/files", summary="上传并解析文件")
async def upload_files(files: List[UploadFile] = File(...)):
    """上传 Word/TXT 文件，后端解析文本后返回，供前端拼接到 Dify query 中。

    返回每个文件的 name 和 content。
    """
    results = []
    for upload_file in files:
        content = ""
        name = upload_file.filename or "上传文件"
        raw = await upload_file.read()

        if name.lower().endswith(".docx"):
            try:
                from docx import Document

                doc = Document(BytesIO(raw))
                parts = []
                for p in doc.paragraphs:
                    if p.text.strip():
                        parts.append(p.text)
                for t in doc.tables:
                    for row in t.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        parts.append(" | ".join(cells))
                content = "\n".join(parts)
            except Exception as exc:
                content = f"(无法解析 .docx 文件：{exc})"
        else:
            try:
                content = raw.decode("utf-8")
            except Exception:
                try:
                    content = raw.decode("gbk")
                except Exception as exc:
                    content = f"(无法读取文件：{exc})"

        results.append({"name": name, "content": content})

    return {"files": results}
